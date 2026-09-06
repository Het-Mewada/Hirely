import pytest
from io import BytesIO
from fastapi import UploadFile
from app.services.parser import ResumeParserService, MANUAL_REVIEW_FLAG
from app.repositories.base import TenantRepository
from app.models.candidate import Candidate
from app.services.storage import StorageService

SAMPLE_PDF_BYTES = b"""%PDF-1.4
1 0 obj <</Type /Catalog /Pages 2 0 R>> endobj
2 0 obj <</Type /Pages /Kids [3 0 R] /Count 1>> endobj
3 0 obj <</Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources <</Font <</F1 4 0 R>>>> /Contents 5 0 R>> endobj
4 0 obj <</Type /Font /Subtype /Type1 /BaseFont /Helvetica>> endobj
5 0 obj <</Length 55>> stream
BT /F1 12 Tf 100 700 Td (Senior Software Engineer Python FastApi) Tj ET
endstream endobj
xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000244 00000 n 
0000000318 00000 n 
trailer <</Size 6 /Root 1 0 R>>
startxref
424
%%EOF
"""

SCANNED_IMAGE_PDF_BYTES = b"""%PDF-1.4
1 0 obj <</Type /Catalog /Pages 2 0 R>> endobj
2 0 obj <</Type /Pages /Kids [3 0 R] /Count 1>> endobj
3 0 obj <</Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R>> endobj
4 0 obj <</Length 10>> stream
% Image Only
endstream endobj
xref
0 5
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000202 00000 n 
trailer <</Size 5 /Root 1 0 R>>
startxref
263
%%EOF
"""

def test_extract_text_from_valid_pdf():
    """
    VERIFIES PDF TEXT EXTRACTION:
    Parses valid text PDF and returns raw extracted string.
    """
    text, needs_review = ResumeParserService.extract_text_from_file("resume.pdf", SAMPLE_PDF_BYTES)
    assert needs_review is False
    assert "Senior Software Engineer" in text
    assert "Python FastApi" in text

def test_extract_text_from_scanned_pdf():
    """
    VERIFIES SCANNED / IMAGE-ONLY PDF DETECTION:
    Returns [NEEDS_MANUAL_REVIEW] flag when text extraction returns < 20 characters.
    """
    text, needs_review = ResumeParserService.extract_text_from_file("scanned_resume.pdf", SCANNED_IMAGE_PDF_BYTES)
    assert needs_review is True
    assert MANUAL_REVIEW_FLAG in text
    assert "Scanned or image-only PDF" in text

def test_extract_text_from_txt_file():
    """
    VERIFIES TXT FILE TEXT EXTRACTION:
    Parses plain text document content correctly.
    """
    txt_content = b"John Doe\nExperienced Fullstack Developer\nSkills: React, Python, PostgreSQL"
    text, needs_review = ResumeParserService.extract_text_from_file("john_resume.txt", txt_content)
    assert needs_review is False
    assert "Experienced Fullstack Developer" in text

def test_candidate_resume_upload_populates_resume_text(db_session, user_org_a_admin, candidate_org_a):
    """
    VERIFIES END-TO-END RESUME TEXT PERSISTENCE IN DB:
    When a candidate uploads a resume, Candidate.resume_text is populated in PostgreSQL.
    """
    repo = TenantRepository(Candidate, db_session, user_org_a_admin.organization_id)
    file = UploadFile(filename="alex_tech_cv.pdf", file=BytesIO(SAMPLE_PDF_BYTES))

    # Perform extraction and update
    rel_url, abs_path = StorageService.save_candidate_resume(
        organization_id=user_org_a_admin.organization_id,
        candidate_id=candidate_org_a.id,
        file=file,
        content=SAMPLE_PDF_BYTES
    )

    extracted_text, needs_review = ResumeParserService.extract_text_from_file("alex_tech_cv.pdf", SAMPLE_PDF_BYTES)

    updated = repo.update(candidate_org_a.id, {
        "resume_url": rel_url,
        "resume_text": extracted_text
    })

    assert updated.resume_text is not None
    assert "Senior Software Engineer" in updated.resume_text
    assert needs_review is False

def test_extract_text_from_docx_file():
    """
    VERIFIES DOCX FILE TEXT EXTRACTION:
    Parses DOCX document paragraphs and tables cleanly.
    """
    import docx

    doc = docx.Document()
    doc.add_paragraph("Het Mewada")
    doc.add_paragraph("Senior Python & Fullstack Developer")
    bio_stream = BytesIO()
    doc.save(bio_stream)
    docx_bytes = bio_stream.getvalue()

    text, needs_review = ResumeParserService.extract_text_from_file("het_resume.docx", docx_bytes)
    assert needs_review is False
    assert "Senior Python & Fullstack Developer" in text

