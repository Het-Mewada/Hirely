import io
import os
import logging
from typing import Tuple

logger = logging.getLogger(__name__)

MANUAL_REVIEW_FLAG = "[NEEDS_MANUAL_REVIEW]"

class ResumeParserService:
    """
    Service for extracting raw text from uploaded resume documents (PDF, DOCX, TXT).
    Detects scanned / image-only PDFs and flags them for manual review.
    """

    @staticmethod
    def extract_text_from_file(filename: str, content: bytes) -> Tuple[str, bool]:
        """
        Extracts raw text from file content based on file extension.

        Returns:
            Tuple[str, bool]: (extracted_text, needs_manual_review_flag)
        """
        _, ext = os.path.splitext(filename.lower())

        try:
            if ext == ".pdf":
                text, review = ResumeParserService._extract_from_pdf(content)
            elif ext == ".docx":
                text, review = ResumeParserService._extract_from_docx(content)
            elif ext == ".txt":
                text, review = ResumeParserService._extract_from_txt(content)
            else:
                text, review = f"{MANUAL_REVIEW_FLAG} Unsupported file format '{ext}' for text extraction.", True

            # Sanitize NUL bytes to avoid PostgreSQL UTF-8 database insertion errors
            clean_text = text.replace("\x00", "")
            return clean_text, review

        except Exception as e:
            logger.error(f"Unhandled error during resume text extraction: {str(e)}")
            return f"{MANUAL_REVIEW_FLAG} Unexpected extraction error. Error: {str(e)}", True

    @staticmethod
    def _extract_from_pdf(content: bytes) -> Tuple[str, bool]:
        """
        Extract text from PDF content using pdfplumber.
        """
        try:
            import pdfplumber

            page_texts = []
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    extracted = page.extract_text()
                    if extracted:
                        page_texts.append(extracted.strip())

            full_text = "\n\n".join(page_texts).strip()

            # Check if PDF is scanned or image-only (minimal or zero text extracted)
            if len(full_text) < 20:
                logger.warning("Scanned or image-only PDF detected during resume extraction.")
                return f"{MANUAL_REVIEW_FLAG} Scanned or image-only PDF. Could not extract raw text.", True

            return full_text, False

        except Exception as e:
            logger.error(f"Failed to parse PDF resume: {str(e)}")
            return f"{MANUAL_REVIEW_FLAG} Failed to parse PDF resume. Error: {str(e)}", True

    @staticmethod
    def _extract_from_docx(content: bytes) -> Tuple[str, bool]:
        """
        Extract text from DOCX document using python-docx.
        """
        try:
            import docx

            doc = docx.Document(io.BytesIO(content))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
            
            # Also extract text from tables inside DOCX
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()])
                    if row_text:
                        paragraphs.append(row_text)

            full_text = "\n".join(paragraphs).strip()

            if len(full_text) < 20:
                return f"{MANUAL_REVIEW_FLAG} DOCX document has minimal or no text content.", True

            return full_text, False

        except Exception as e:
            logger.error(f"Failed to parse DOCX resume: {str(e)}")
            return f"{MANUAL_REVIEW_FLAG} Failed to parse DOCX document. Error: {str(e)}", True

    @staticmethod
    def _extract_from_txt(content: bytes) -> Tuple[str, bool]:
        """
        Extract text from TXT file content.
        """
        try:
            text = content.decode("utf-8", errors="ignore").strip()
            if len(text) < 10:
                return f"{MANUAL_REVIEW_FLAG} Empty or unreadable text file.", True
            return text, False
        except Exception as e:
            return f"{MANUAL_REVIEW_FLAG} Failed to read text file. Error: {str(e)}", True
